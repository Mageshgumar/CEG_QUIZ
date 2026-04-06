from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CEG Quiz Automation System\n")
    run.bold = True
    run.font.size = 240000  # ~24 pt

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Project Documentation\n")

    doc.add_paragraph("\n")

    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run("Prepared By: ____________________\n")
    details.add_run("Department: ____________________\n")
    details.add_run("Institution: ____________________\n")
    details.add_run(f"Date: {date.today().isoformat()}\n")

    doc.add_page_break()


def add_section(doc: Document, heading: str, points: list[str]) -> None:
    doc.add_heading(heading, level=1)
    for p in points:
        doc.add_paragraph(p)


def main() -> None:
    doc = Document()

    add_title_page(doc)

    add_section(
        doc,
        "Abstract",
        [
            "CEG Quiz Automation System is a Telegram-based assessment platform with an integrated teacher dashboard for managing tests and reviewing student performance.",
            "The system supports voice and text answers, timed questions, random question selection, one-time test eligibility by roll number, parent notifications, and secure teacher test operations (create/edit/delete/activate).",
            "The platform improves assessment efficiency and provides structured analytics while reducing manual correction and reporting work.",
        ],
    )

    add_section(
        doc,
        "Introduction",
        [
            "Conventional classroom quizzes often require manual preparation, invigilation, evaluation, and communication of results.",
            "This project automates the full quiz lifecycle using a Telegram bot for students and a web dashboard for teachers.",
            "Students complete registration, choose active tests, answer questions (button/text/voice), and receive immediate feedback. Teachers configure assessments and monitor attempts with detailed breakdowns.",
        ],
    )

    add_section(
        doc,
        "Project Description",
        [
            "Core Modules:",
            "1. Telegram Bot Module: Handles student registration, test selection, timed MCQ flow, answer validation, voice transcription, result summary, and parent alerts.",
            "2. Teacher Dashboard Module: Login-protected interface for creating/editing/deleting tests, setting active/inactive status, sorting/searching attempts, and viewing detailed student responses.",
            "3. Data Persistence Module: Stores tests, attempts, parent mappings, and results in JSON files for lightweight persistence.",
            "4. Validation and Access Module: Enforces 10-digit phone/roll constraints and one-time eligibility rules.",
        ],
    )

    add_section(
        doc,
        "AR Diagram (Architecture Diagram)",
        [
            "High-Level Architecture:",
            "Teacher -> Dashboard (Flask) -> Test Store (tests.json)",
            "Student -> Telegram Bot (python-telegram-bot) -> Quiz Engine -> Attempts Store (attempts.json)",
            "Bot -> Parent Notification -> Telegram Parent Chat",
            "Bot <-> Speech Recognition + ffmpeg pipeline for voice decoding",
            "Dashboard <-> Attempts Data for reporting and filtering",
        ],
    )

    add_section(
        doc,
        "Project Flow",
        [
            "1. Teacher logs in to dashboard.",
            "2. Teacher creates or updates test (question file upload, timer, random count, one-time option, active status).",
            "3. Student sends /start in Telegram bot and submits name, phone (10 digits), roll (10 digits), and parent username.",
            "4. Student chooses one of the currently active tests.",
            "5. Bot serves timed questions; student answers via button/text/voice.",
            "6. Bot evaluates each answer, applies timeout logic, and moves to next question.",
            "7. Final score and detailed breakdown are shown to student and sent to parent.",
            "8. Teacher views attempts, sorts/searches records, checks detailed answers, or deletes attempts if needed.",
        ],
    )

    add_section(
        doc,
        "Inputs and Outputs",
        [
            "Inputs:",
            "- Teacher credentials (username/password)",
            "- Test metadata (name, timer, random count, one-time setting, active flag)",
            "- Question text file in MCQ format (Q, A/B/C/D, ANS)",
            "- Student details (name, 10-digit phone, 10-digit roll, parent username)",
            "- Student answer signals (button/text/voice)",
            "Outputs:",
            "- Real-time question prompts and correctness feedback",
            "- Time-out handling and question progression",
            "- Student result summary with correct/wrong details",
            "- Parent result notification in Telegram",
            "- Dashboard attempt tables, sort/search views, and per-question audit data",
        ],
    )

    add_section(
        doc,
        "Conclusion",
        [
            "The CEG Quiz Automation System successfully combines conversational assessment through Telegram with teacher-centric administrative control via dashboard.",
            "It ensures structured test execution, robust input validation, flexible assessment controls, and transparent result tracking.",
            "Future enhancements can include database migration, role-based accounts, richer analytics, and cloud deployment for institutional scaling.",
        ],
    )

    out_path = "Project_Documentation.docx"
    doc.save(out_path)
    print(f"Created {out_path}")


if __name__ == "__main__":
    main()
